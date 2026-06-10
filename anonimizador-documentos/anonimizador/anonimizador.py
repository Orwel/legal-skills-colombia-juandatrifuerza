#!/usr/bin/env python3
"""
Anonimizador local de documentos jurídicos (Colombia) — v6
==========================================================
Detección automática de datos personales + revisión antes de guardar.
El dato real nunca sale de tu máquina.

  A) Interfaz gráfica:  python anonimizador.py
  B) Línea de comandos: python anonimizador.py expediente.docx

GENERA:
    expediente_ANONIMIZADO.docx   -> lo único que subes a la IA
    expediente_EQUIVALENCIAS.csv  -> se queda en tu disco (NUNCA lo subas)
"""

import csv
import json
import os
import re
import shutil
import sys
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path
from xml.sax.saxutils import escape, unescape

from docx.oxml.ns import qn

from motor_anonimizacion import Hallazgo, MotorAnonimizacion

if getattr(sys, "frozen", False):
    BASE_DIR = Path(sys.executable).parent
else:
    BASE_DIR = Path(__file__).parent

ARCHIVO_REEMPLAZOS = BASE_DIR / "reemplazos.json"
REEMPLAZOS_POR_DEFECTO: dict = {}

ADVERTENCIA = (
    "Revisa el documento anonimizado antes de subirlo. Aunque el sistema detecta "
    "automáticamente nombres, cédulas y empresas, la revisión final es criterio del abogado."
)


def cargar_reemplazos() -> dict:
    if ARCHIVO_REEMPLAZOS.exists():
        try:
            return json.loads(ARCHIVO_REEMPLAZOS.read_text(encoding="utf-8"))
        except Exception:
            pass
    return dict(REEMPLAZOS_POR_DEFECTO)


def guardar_reemplazos(reemplazos: dict):
    ARCHIVO_REEMPLAZOS.write_text(
        json.dumps(reemplazos, ensure_ascii=False, indent=2), encoding="utf-8")


def _extraer_texto_contenedor(contenedor) -> list[str]:
    partes = []
    for p in contenedor.paragraphs:
        if p.text.strip():
            partes.append(p.text)
    for tabla in contenedor.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                partes.extend(_extraer_texto_contenedor(celda))
    return partes


def extraer_texto_docx(ruta: Path) -> str:
    from docx import Document
    doc = Document(ruta)
    partes = _extraer_texto_contenedor(doc)
    for seccion in doc.sections:
        for parte in (seccion.header, seccion.footer,
                      seccion.first_page_header, seccion.first_page_footer,
                      seccion.even_page_header, seccion.even_page_footer):
            if parte is not None:
                partes.extend(_extraer_texto_contenedor(parte))
    return "\n".join(partes)


def extraer_parrafos(ruta: Path) -> list[str]:
    if ruta.suffix.lower() == ".docx":
        from docx import Document
        doc = Document(ruta)
        return [p.text for p in doc.paragraphs if p.text.strip()]
    return [ln for ln in ruta.read_text(encoding="utf-8", errors="ignore").splitlines() if ln.strip()]


def extraer_texto(ruta: Path) -> str:
    if ruta.suffix.lower() == ".docx":
        return extraer_texto_docx(ruta)
    return ruta.read_text(encoding="utf-8", errors="ignore")


def _copiar_rpr(run) -> object | None:
    rpr = run._element.find(qn("w:rPr"))
    return deepcopy(rpr) if rpr is not None else None


def _aplicar_rpr(run, rpr) -> None:
    if rpr is None:
        return
    existente = run._element.find(qn("w:rPr"))
    if existente is not None:
        run._element.remove(existente)
    run._element.insert(0, deepcopy(rpr))


def _mapear_caracteres_runs(p) -> list[int]:
    """Para cada carácter del párrafo, índice del run que lo contiene."""
    mapa: list[int] = []
    for i, run in enumerate(p.runs):
        mapa.extend([i] * len(run.text))
    return mapa


def _segmentos_con_formato(texto: str, mapa: list[int],
                           spans: list[tuple[int, int, str, str]]) -> list[tuple[str, int]]:
    if not spans:
        return [(texto, mapa[0] if mapa else 0)]
    segmentos: list[tuple[str, int]] = []
    pos = 0
    for inicio, fin, reemplazo, _ in sorted(spans, key=lambda s: s[0]):
        if pos < inicio:
            fmt = mapa[pos] if pos < len(mapa) else (mapa[-1] if mapa else 0)
            segmentos.append((texto[pos:inicio], fmt))
        fmt_rep = mapa[inicio] if inicio < len(mapa) else (mapa[-1] if mapa else 0)
        segmentos.append((reemplazo, fmt_rep))
        pos = fin
    if pos < len(texto):
        fmt = mapa[pos] if pos < len(mapa) else (mapa[-1] if mapa else 0)
        segmentos.append((texto[pos:], fmt))
    return segmentos


def _fusionar_segmentos(segmentos: list[tuple[str, int]]) -> list[tuple[str, int]]:
    if not segmentos:
        return []
    fusionados = [segmentos[0]]
    for texto, fmt in segmentos[1:]:
        if not texto:
            continue
        if fmt == fusionados[-1][1]:
            fusionados[-1] = (fusionados[-1][0] + texto, fmt)
        else:
            fusionados.append((texto, fmt))
    return fusionados


def _eliminar_runs(p) -> None:
    for run in list(p.runs):
        run._element.getparent().remove(run._element)


def _aplicar_a_parrafo(p, motor: MotorAnonimizacion, hallazgos: list[Hallazgo],
                       registro: dict):
    """Aplica reemplazos conservando el formato (subrayado, negrita, etc.) de cada run."""
    texto_original = p.text
    if not texto_original:
        return

    spans = motor._recolectar_spans(texto_original, hallazgos)
    if not spans:
        return

    mapa = _mapear_caracteres_runs(p)
    rpr_copias = [_copiar_rpr(run) for run in p.runs]
    segmentos = _fusionar_segmentos(_segmentos_con_formato(texto_original, mapa, spans))

    _eliminar_runs(p)
    for texto, fmt_idx in segmentos:
        if not texto:
            continue
        nuevo_run = p.add_run(texto)
        if rpr_copias and fmt_idx < len(rpr_copias):
            _aplicar_rpr(nuevo_run, rpr_copias[fmt_idx])

    for _, _, reemplazo, original in spans:
        registro[original] = reemplazo


def _procesar_contenedor(contenedor, motor: MotorAnonimizacion,
                         hallazgos: list[Hallazgo], registro: dict):
    for p in contenedor.paragraphs:
        _aplicar_a_parrafo(p, motor, hallazgos, registro)
    for tabla in contenedor.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                _procesar_contenedor(celda, motor, hallazgos, registro)


def _anonimizar_xml_wt(xml_texto: str, motor: MotorAnonimizacion,
                       hallazgos: list[Hallazgo], registro: dict) -> str:
    def _sub(m):
        contenido = unescape(m.group(2))
        nuevo = motor.aplicar(contenido, hallazgos, registro)
        return m.group(1) + escape(nuevo) + m.group(3)

    return re.sub(r"(<w:t[^>]*>)(.*?)(</w:t>)", _sub, xml_texto, flags=re.DOTALL)


def _postprocesar_zip(ruta_docx: Path, motor: MotorAnonimizacion,
                      hallazgos: list[Hallazgo], registro: dict):
    objetivos = {"word/footnotes.xml", "word/endnotes.xml", "word/comments.xml"}
    fd, tmp_nombre = tempfile.mkstemp(suffix=".docx")
    tmp = Path(tmp_nombre)
    os.close(fd)
    with zipfile.ZipFile(ruta_docx, "r") as zin, \
         zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            datos = zin.read(item.filename)
            if item.filename in objetivos:
                xml = datos.decode("utf-8")
                datos = _anonimizar_xml_wt(xml, motor, hallazgos, registro).encode("utf-8")
            zout.writestr(item, datos)
    shutil.move(str(tmp), str(ruta_docx))


def procesar_docx(ruta: Path, motor: MotorAnonimizacion,
                  hallazgos: list[Hallazgo], registro: dict) -> Path:
    from docx import Document
    doc = Document(ruta)
    _procesar_contenedor(doc, motor, hallazgos, registro)
    for seccion in doc.sections:
        for parte in (seccion.header, seccion.footer,
                      seccion.first_page_header, seccion.first_page_footer,
                      seccion.even_page_header, seccion.even_page_footer):
            if parte is not None:
                _procesar_contenedor(parte, motor, hallazgos, registro)
    cp = doc.core_properties
    cp.author = cp.last_modified_by = cp.title = cp.subject = ""
    cp.comments = cp.keywords = cp.category = ""
    salida = ruta.with_name(f"{ruta.stem}_ANONIMIZADO.docx")
    doc.save(salida)
    _postprocesar_zip(salida, motor, hallazgos, registro)
    return salida


def procesar_txt(ruta: Path, motor: MotorAnonimizacion,
                 hallazgos: list[Hallazgo], registro: dict) -> Path:
    texto = ruta.read_text(encoding="utf-8", errors="ignore")
    salida = ruta.with_name(f"{ruta.stem}_ANONIMIZADO{ruta.suffix}")
    salida.write_text(motor.aplicar(texto, hallazgos, registro), encoding="utf-8")
    return salida


def guardar_equivalencias(ruta: Path, registro: dict) -> Path:
    tabla = ruta.with_name(f"{ruta.stem}_EQUIVALENCIAS.csv")
    with open(tabla, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.writer(f)
        w.writerow(["Dato real", "Reemplazo", "Tipo"])
        for real, rol in sorted(registro.items(), key=lambda x: x[0].lower()):
            w.writerow([real, rol, ""])
    return tabla


def procesar(ruta: Path, hallazgos: list[Hallazgo]):
    motor = MotorAnonimizacion()
    registro: dict[str, str] = {}
    if ruta.suffix.lower() == ".docx":
        salida = procesar_docx(ruta, motor, hallazgos, registro)
    elif ruta.suffix.lower() in (".txt", ".md"):
        salida = procesar_txt(ruta, motor, hallazgos, registro)
    else:
        raise ValueError("Formato no soportado. Usa .docx, .txt o .md")
    tabla = guardar_equivalencias(ruta, registro)
    parrafos_salida = extraer_parrafos(salida)
    fugas = motor.verificar_resultado(parrafos_salida, hallazgos)
    return salida, tabla, len(registro), fugas


def _parsear_manual(texto_caja: str) -> dict:
    nuevos = {}
    for linea in texto_caja.splitlines():
        linea = linea.strip()
        if not linea or linea.startswith("#"):
            continue
        if "=>" in linea:
            real, rol = linea.split("=>", 1)
            if real.strip() and rol.strip():
                nuevos[real.strip()] = rol.strip()
    return nuevos


def _crear_tabla_revision(padre, items: list[Hallazgo], titulo_cols: tuple,
                          valores_defecto: bool):
    import tkinter as tk
    from tkinter import ttk

    marco = tk.Frame(padre, bg="#111111")
    marco.pack(fill="both", expand=True)

    canvas = tk.Canvas(marco, bg="#1d1d1d", highlightthickness=0)
    scroll = ttk.Scrollbar(marco, orient="vertical", command=canvas.yview)
    interior = tk.Frame(canvas, bg="#1d1d1d")
    interior.bind("<Configure>",
                  lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=interior, anchor="nw")
    canvas.configure(yscrollcommand=scroll.set)
    canvas.pack(side="left", fill="both", expand=True)
    scroll.pack(side="right", fill="y")

    vars_check: list[tk.BooleanVar] = []
    colores_tipo = {
        "PERSONA": "#7ec8ff", "EMPRESA": "#ffb347", "LUGAR": "#b19cd9",
        "CÉDULA": "#ff6b6b", "NIT": "#ff6b6b", "CORREO": "#ff6b6b",
        "TELÉFONO": "#ff6b6b", "DIRECCIÓN": "#ff6b6b", "RADICADO": "#aaaaaa",
        "MANUAL": "#9be29b", "REVISAR": "#ffcc66",
    }

    encabezado = tk.Frame(interior, bg="#2a2a2a")
    encabezado.pack(fill="x", pady=(0, 4))
    for i, (txt, w) in enumerate(titulo_cols):
        tk.Label(encabezado, text=txt, bg="#2a2a2a", fg="#cccccc",
                 font=("Segoe UI", 9, "bold"), width=w, anchor="w"
                 ).grid(row=0, column=i, padx=4, pady=4, sticky="w")

    for h in items:
        fila = tk.Frame(interior, bg="#1d1d1d")
        fila.pack(fill="x", pady=1)
        var = tk.BooleanVar(value=valores_defecto if h.tipo == "REVISAR" else h.activo)
        vars_check.append(var)
        color = colores_tipo.get(h.tipo, "#f5f5f5")
        tk.Checkbutton(fila, variable=var, bg="#1d1d1d", activebackground="#1d1d1d",
                       selectcolor="#333333").grid(row=0, column=0, padx=4)
        tk.Label(fila, text=h.texto[:45] + ("…" if len(h.texto) > 45 else ""),
                 bg="#1d1d1d", fg="#f5f5f5", font=("Consolas", 9),
                 width=30, anchor="w").grid(row=0, column=1, sticky="w")
        tk.Label(fila, text=h.reemplazo, bg="#1d1d1d", fg="#9be29b",
                 font=("Consolas", 9), width=18, anchor="w").grid(row=0, column=2, sticky="w")
        tk.Label(fila, text=h.tipo, bg="#1d1d1d", fg=color,
                 font=("Segoe UI", 8), width=10, anchor="w").grid(row=0, column=3, sticky="w")
        tk.Label(fila, text=h.fuente, bg="#1d1d1d", fg="#888888",
                 font=("Segoe UI", 8), width=10, anchor="w").grid(row=0, column=4, sticky="w")
        conf_color = {"alta": "#9be29b", "media": "#ffcc66", "baja": "#ff6b6b"}.get(h.confianza, "#888")
        tk.Label(fila, text=h.confianza, bg="#1d1d1d", fg=conf_color,
                 font=("Segoe UI", 8), width=8, anchor="w").grid(row=0, column=5, sticky="w")

    return vars_check


def mostrar_revision(raiz, hallazgos: list[Hallazgo], texto: str,
                     motor: MotorAnonimizacion,
                     parrafos: list[str] | None = None) -> list[Hallazgo] | None:
    import tkinter as tk
    from tkinter import ttk

    resultado: list[Hallazgo] | None = None
    bloques = parrafos if parrafos else texto.split("\n")
    pendientes = motor.detectar_pendientes(texto, hallazgos, parrafos=bloques)

    ventana = tk.Toplevel(raiz)
    ventana.title("Revisión — confirma qué anonimizar")
    ventana.geometry("780x560")
    ventana.configure(bg="#111111")
    ventana.transient(raiz)
    ventana.grab_set()

    tk.Label(
        ventana,
        text=(f"Detectados: {len(hallazgos)} · Pendientes por revisar: {len(pendientes)}. "
              "Revisa la pestaña «Pendientes» antes de confirmar."),
        bg="#111111", fg="#f5f5f5", font=("Segoe UI", 10), wraplength=720,
    ).pack(pady=(12, 8), padx=16)

    notebook = ttk.Notebook(ventana)
    notebook.pack(fill="both", expand=True, padx=16, pady=4)

    tab_detectados = tk.Frame(notebook, bg="#111111")
    tab_pendientes = tk.Frame(notebook, bg="#111111")
    notebook.add(tab_detectados, text=f"Detectados ({len(hallazgos)})")
    notebook.add(tab_pendientes, text=f"Pendientes ({len(pendientes)})")

    cols = [("✓", 3), ("Dato", 30), ("Reemplazo", 18), ("Tipo", 10), ("Fuente", 10), ("Conf.", 8)]
    vars_hallazgos = _crear_tabla_revision(tab_detectados, hallazgos, cols, True)
    vars_pendientes = _crear_tabla_revision(tab_pendientes, pendientes, cols, False) if pendientes else []

    if not pendientes:
        tk.Label(tab_pendientes, text="No quedaron números ni nombres sospechosos sin anonimizar.",
                 bg="#111111", fg="#9be29b", font=("Segoe UI", 10)).pack(pady=40)

    def confirmar():
        nonlocal resultado
        for h, var in zip(hallazgos, vars_hallazgos):
            h.activo = var.get()
        extras: list[Hallazgo] = []
        for h, var in zip(pendientes, vars_pendientes):
            if var.get():
                h.activo = True
                h.tipo = "MANUAL"
                h.fuente = "revisión"
                h.patron = re.compile(r"(?<!\w)" + re.escape(h.texto) + r"(?!\w)", re.IGNORECASE)
                extras.append(h)
        resultado = hallazgos + extras
        ventana.destroy()

    def cancelar():
        ventana.destroy()

    marco_btn = tk.Frame(ventana, bg="#111111")
    marco_btn.pack(pady=12)
    tk.Button(marco_btn, text="Confirmar y anonimizar", command=confirmar,
              bg="#ffffff", fg="#111111", font=("Segoe UI", 10, "bold"),
              relief="flat", padx=16, pady=6).pack(side="left", padx=8)
    tk.Button(marco_btn, text="Cancelar", command=cancelar,
              bg="#333333", fg="#f5f5f5", font=("Segoe UI", 10),
              relief="flat", padx=16, pady=6).pack(side="left", padx=8)

    ventana.wait_window()
    return resultado


def lanzar_gui():
    import tkinter as tk
    from tkinter import filedialog, messagebox, scrolledtext

    reemplazos = cargar_reemplazos()
    archivo_sel: dict = {"ruta": None}
    hallazgos_actuales: list[Hallazgo] = []
    parrafos_actual: dict = {"lista": []}
    texto_actual: dict = {"texto": ""}
    motor_actual: dict = {"motor": None}

    raiz = tk.Tk()
    raiz.title("Anonimizador de Expedientes — Trifuerza v6")
    raiz.geometry("700x640")
    raiz.configure(bg="#111111")
    estilo = {"bg": "#111111", "fg": "#f5f5f5", "font": ("Segoe UI", 10)}

    tk.Label(raiz, text="ANONIMIZADOR DE EXPEDIENTES",
             font=("Segoe UI", 14, "bold"), bg="#111111", fg="#ffffff"
             ).pack(pady=(16, 2))
    tk.Label(raiz, text="Detección automática · Revisión antes de guardar · Todo local",
             **estilo).pack(pady=(0, 12))

    marco_archivo = tk.Frame(raiz, bg="#111111")
    marco_archivo.pack(fill="x", padx=20)
    etiqueta_archivo = tk.Label(marco_archivo, text="Ningún archivo seleccionado",
                                anchor="w", **estilo)

    def seleccionar():
        ruta = filedialog.askopenfilename(
            title="Selecciona el documento",
            filetypes=[("Documentos", "*.docx *.txt *.md")])
        if ruta:
            archivo_sel["ruta"] = Path(ruta)
            etiqueta_archivo.config(text=Path(ruta).name)
            hallazgos_actuales.clear()
            log("Archivo seleccionado. Pulsa ANALIZAR para detectar datos sensibles.")

    tk.Button(marco_archivo, text="Seleccionar documento…", command=seleccionar,
              bg="#ffffff", fg="#111111", font=("Segoe UI", 10, "bold"),
              relief="flat", padx=12, pady=4).pack(side="left")
    etiqueta_archivo.pack(side="left", padx=12)

    tk.Label(raiz, text="Reemplazos adicionales (opcional, formato:  Nombre => rol):",
             anchor="w", **estilo).pack(fill="x", padx=20, pady=(12, 2))
    tk.Label(raiz,
             text="Cédulas, NIT, nombres y empresas se detectan solos. "
                  "Cortes, leyes y juzgados NO se anonimizan.",
             anchor="w", fg="#aaaaaa", bg="#111111", font=("Segoe UI", 9)
             ).pack(fill="x", padx=20, pady=(0, 4))

    caja = scrolledtext.ScrolledText(raiz, height=5, font=("Consolas", 10),
                                     bg="#1d1d1d", fg="#f5f5f5",
                                     insertbackground="#ffffff", relief="flat")
    caja.pack(fill="both", expand=False, padx=20)
    if reemplazos:
        caja.insert("1.0", "\n".join(f"{k} => {v}" for k, v in reemplazos.items()))

    salida_log = scrolledtext.ScrolledText(raiz, height=10, font=("Consolas", 9),
                                           bg="#1d1d1d", fg="#9be29b",
                                           relief="flat", state="disabled")

    def log(msg):
        salida_log.config(state="normal")
        salida_log.insert("end", msg + "\n")
        salida_log.config(state="disabled")
        salida_log.see("end")

    def analizar():
        if not archivo_sel["ruta"]:
            messagebox.showwarning("Falta el archivo", "Primero selecciona un documento.")
            return
        manual = _parsear_manual(caja.get("1.0", "end"))
        guardar_reemplazos(manual)
        try:
            texto = extraer_texto(archivo_sel["ruta"])
            parrafos_actual["lista"] = extraer_parrafos(archivo_sel["ruta"])
            motor = MotorAnonimizacion(reemplazos_manual=manual)
            hallazgos = motor.analizar(texto)
            texto_actual["texto"] = texto
            motor_actual["motor"] = motor
            hallazgos_actuales.clear()
            hallazgos_actuales.extend(hallazgos)
            log(f"Análisis completado: {len(hallazgos)} hallazgos.")
            for h in hallazgos[:8]:
                log(f"  · {h.texto[:40]} → {h.reemplazo}  ({h.tipo})")
            if len(hallazgos) > 8:
                log(f"  … y {len(hallazgos) - 8} más.")
            if not hallazgos:
                log("  No se detectaron datos sensibles. Revisa el documento manualmente.")
            else:
                log("Pulsa ANONIMIZAR para revisar y confirmar.")
        except ImportError as e:
            messagebox.showerror("Falta una librería", str(e))
        except Exception as e:
            messagebox.showerror("Error", str(e))

    def anonimizar():
        if not archivo_sel["ruta"]:
            messagebox.showwarning("Falta el archivo", "Primero selecciona un documento.")
            return
        if not hallazgos_actuales:
            analizar()
            if not hallazgos_actuales:
                return
        if not texto_actual["texto"] or motor_actual["motor"] is None:
            analizar()
        confirmados = mostrar_revision(
            raiz, list(hallazgos_actuales), texto_actual["texto"], motor_actual["motor"],
            parrafos=parrafos_actual["lista"])
        if not confirmados:
            log("Anonimización cancelada.")
            return
        activos = [h for h in confirmados if h.activo]
        if not activos:
            messagebox.showwarning("Sin cambios", "No hay datos marcados para anonimizar.")
            return
        try:
            salida, tabla, n, fugas = procesar(archivo_sel["ruta"], confirmados)
            log(f"✔ Documento anonimizado: {salida.name}")
            log(f"✔ Equivalencias: {tabla.name}  (NUNCA la subas)")
            log(f"  Reemplazos aplicados: {n}")
            if fugas:
                log(f"⚠ ATENCIÓN: quedaron {len(fugas)} dato(s) sin anonimizar:")
                for f in fugas[:5]:
                    log(f"    · {f.texto} ({f.fuente})")
            log(f"⚠ {ADVERTENCIA}")
            msg = f"Documento guardado en:\n{salida}\n\nEquivalencias (no subir):\n{tabla}"
            if fugas:
                msg += f"\n\n⚠ Quedaron {len(fugas)} dato(s) visibles. Revisa el documento."
            messagebox.showinfo("Listo", msg)
        except Exception as e:
            messagebox.showerror("Error", str(e))

    marco_btn = tk.Frame(raiz, bg="#111111")
    marco_btn.pack(pady=10)
    tk.Button(marco_btn, text="ANALIZAR", command=analizar,
              bg="#333333", fg="#f5f5f5", font=("Segoe UI", 10, "bold"),
              relief="flat", padx=20, pady=6).pack(side="left", padx=6)
    tk.Button(marco_btn, text="ANONIMIZAR", command=anonimizar,
              bg="#ffffff", fg="#111111", font=("Segoe UI", 11, "bold"),
              relief="flat", padx=24, pady=8).pack(side="left", padx=6)

    salida_log.pack(fill="both", expand=True, padx=20, pady=(0, 16))
    log("Selecciona un documento y pulsa ANALIZAR.")
    raiz.mainloop()


def main():
    if len(sys.argv) >= 2:
        ruta = Path(sys.argv[1])
        if not ruta.exists():
            print(f"No existe el archivo: {ruta}")
            sys.exit(1)
        manual = cargar_reemplazos()
        texto = extraer_texto(ruta)
        motor = MotorAnonimizacion(reemplazos_manual=manual)
        hallazgos = motor.analizar(texto)
        salida, tabla, n, fugas = procesar(ruta, hallazgos)
        print(f"✔ Documento anonimizado: {salida.name}")
        print(f"✔ Equivalencias: {tabla.name}  (NUNCA subas este archivo)")
        print(f"  Reemplazos aplicados: {n}")
        if fugas:
            print(f"⚠ Quedaron {len(fugas)} dato(s) sin anonimizar")
        print(f"⚠ {ADVERTENCIA}")
    else:
        lanzar_gui()


if __name__ == "__main__":
    main()
