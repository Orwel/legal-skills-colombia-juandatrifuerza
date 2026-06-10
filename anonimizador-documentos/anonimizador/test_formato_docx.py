"""Prueba: el anonimizador no debe extender subrayado a todo el párrafo."""

import tempfile
from pathlib import Path

from docx import Document
from docx.enum.text import WD_UNDERLINE

from anonimizador import _aplicar_a_parrafo, _procesar_contenedor
from motor_anonimizacion import MotorAnonimizacion


def _tiene_subrayado(run) -> bool:
    u = run.underline
    return u is not None and u is not False and u != WD_UNDERLINE.NONE


def test_no_extiende_subrayado_a_todo_parrafo():
    doc = Document()
    p = doc.add_paragraph()
    p.add_run("La solicitud vincula a ")
    sub = p.add_run("ACME DEL VALLE S.A.S.")
    sub.underline = True
    p.add_run(" como tercera responsable del daño.")

    motor = MotorAnonimizacion()
    texto = p.text
    hallazgos = motor.analizar(texto)
    registro = {}
    _aplicar_a_parrafo(p, motor, hallazgos, registro)

    runs_sub = [r for r in p.runs if _tiene_subrayado(r)]
    runs_sin = [r for r in p.runs if not _tiene_subrayado(r)]
    assert any("[EMPRESA" in r.text for r in p.runs)
    assert runs_sin, "Debe haber runs sin subrayado"
    assert len(runs_sin) >= 1
    assert not all(_tiene_subrayado(r) for r in p.runs), "No todo el párrafo puede quedar subrayado"
    assert "solicitud" in p.text and "tercera" in p.text


def test_docx_guardado_conserva_formato():
    with tempfile.TemporaryDirectory() as tmp:
        ruta = Path(tmp) / "prueba.docx"
        doc = Document()
        p = doc.add_paragraph()
        p.add_run("Punto 5. Texto normal ")
        u = p.add_run("SERVIGAS DEL NORTE S.A. E.S.P.")
        u.underline = True
        p.add_run(" y más texto normal.")
        doc.save(ruta)

        motor = MotorAnonimizacion()
        hallazgos = motor.analizar(p.text)
        registro = {}
        doc2 = Document(ruta)
        _procesar_contenedor(doc2, motor, hallazgos, registro)
        salida = ruta.with_name("salida.docx")
        doc2.save(salida)

        ver = Document(salida)
        para = ver.paragraphs[0]
        subrayados = sum(1 for r in para.runs if _tiene_subrayado(r))
        assert subrayados <= 2
        assert any("[EMPRESA" in r.text for r in para.runs)
        assert any(not _tiene_subrayado(r) and r.text.strip() for r in para.runs)


if __name__ == "__main__":
    test_no_extiende_subrayado_a_todo_parrafo()
    print("OK  test_no_extiende_subrayado_a_todo_parrafo")
    test_docx_guardado_conserva_formato()
    print("OK  test_docx_guardado_conserva_formato")
