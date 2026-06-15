#!/usr/bin/env python3
"""
Anonimizador local de documentos jurídicos (Colombia) — v7
==========================================================
Detección automática de datos personales + revisión antes de guardar.
El dato real nunca sale de tu máquina.

  A) Interfaz gráfica:  python anonimizador.py
  B) Línea de comandos: python anonimizador.py doc1.docx --salida C:/Salida/

ENTRADA:  .docx · .txt · .md · .pdf (texto seleccionable, no escaneados)
SALIDA:   mismo formato + _ANONIMIZADO  y  _EQUIVALENCIAS.csv (no subir)
"""

import csv
import json
import os
import re
import shutil
import sys
import tempfile
import zipfile
from copy import deepcopy
from pathlib import Path
from xml.sax.saxutils import escape, unescape

from docx.oxml.ns import qn

from motor_anonimizacion import Hallazgo, MotorAnonimizacion
from procesar_pdf import extraer_parrafos_pdf, extraer_texto_pdf, procesar_pdf
from rutas_app import cargar_preferencias, guardar_preferencias, ruta_reemplazos

EXTENSIONES_SOPORTADAS = {".docx", ".txt", ".md", ".pdf"}
FILTRO_ARCHIVOS = [
    ("Documentos", "*.docx *.txt *.md *.pdf"),
    ("Word", "*.docx"),
    ("PDF", "*.pdf"),
    ("Texto", "*.txt *.md"),
    ("Todos", "*.*"),
]

ARCHIVO_REEMPLAZOS = ruta_reemplazos()
REEMPLAZOS_POR_DEFECTO: dict = {}

ADVERTENCIA = (
    "Revisa el documento anonimizado antes de subirlo. Aunque el sistema detecta "
    "automáticamente nombres, cédulas y empresas, la revisión final es criterio del abogado."
)


def cargar_reemplazos() -> dict:
    if ARCHIVO_REEMPLAZOS.exists():
        try:
            return json.loads(ARCHIVO_REEMPLAZOS.read_text(encoding="utf-8"))
        except Exception:
            pass
    return dict(REEMPLAZOS_POR_DEFECTO)


def guardar_reemplazos(reemplazos: dict):
    ARCHIVO_REEMPLAZOS.write_text(
        json.dumps(reemplazos, ensure_ascii=False, indent=2), encoding="utf-8")


def _extraer_texto_contenedor(contenedor) -> list[str]:
    partes = []
    for p in contenedor.paragraphs:
        if p.text.strip():
            partes.append(p.text)
    for tabla in contenedor.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                partes.extend(_extraer_texto_contenedor(celda))
    return partes


def extraer_texto_docx(ruta: Path) -> str:
    from docx import Document
    doc = Document(ruta)
    partes = _extraer_texto_contenedor(doc)
    for seccion in doc.sections:
        for parte in (seccion.header, seccion.footer,
                      seccion.first_page_header, seccion.first_page_footer,
                      seccion.even_page_header, seccion.even_page_footer):
            if parte is not None:
                partes.extend(_extraer_texto_contenedor(parte))
    return "\n".join(partes)


def es_formato_soportado(ruta: Path) -> bool:
    return ruta.suffix.lower() in EXTENSIONES_SOPORTADAS


def recolectar_rutas(entradas: list[str]) -> list[Path]:
    rutas: list[Path] = []
    for entrada in entradas:
        p = Path(entrada)
        if not p.exists():
            raise FileNotFoundError(f"No existe: {p}")
        if p.is_dir():
            for f in sorted(p.iterdir()):
                if f.is_file() and es_formato_soportado(f):
                    rutas.append(f.resolve())
        elif es_formato_soportado(p):
            rutas.append(p.resolve())
        else:
            raise ValueError(f"Formato no soportado: {p.name}")
    return rutas


def extraer_parrafos(ruta: Path) -> list[str]:
    suf = ruta.suffix.lower()
    if suf == ".docx":
        from docx import Document
        doc = Document(ruta)
        return [p.text for p in doc.paragraphs if p.text.strip()]
    if suf == ".pdf":
        return extraer_parrafos_pdf(ruta)
    return [ln for ln in ruta.read_text(encoding="utf-8", errors="ignore").splitlines() if ln.strip()]


def extraer_texto(ruta: Path) -> str:
    suf = ruta.suffix.lower()
    if suf == ".docx":
        return extraer_texto_docx(ruta)
    if suf == ".pdf":
        return extraer_texto_pdf(ruta)
    return ruta.read_text(encoding="utf-8", errors="ignore")


def _copiar_rpr(run) -> object | None:
    rpr = run._element.find(qn("w:rPr"))
    return deepcopy(rpr) if rpr is not None else None


def _aplicar_rpr(run, rpr) -> None:
    if rpr is None:
        return
    existente = run._element.find(qn("w:rPr"))
    if existente is not None:
        run._element.remove(existente)
    run._element.insert(0, deepcopy(rpr))


def _mapear_caracteres_runs(p) -> list[int]:
    """Para cada carácter del párrafo, índice del run que lo contiene."""
    mapa: list[int] = []
    for i, run in enumerate(p.runs):
        mapa.extend([i] * len(run.text))
    return mapa


def _segmentos_con_formato(texto: str, mapa: list[int],
                           spans: list[tuple[int, int, str, str]]) -> list[tuple[str, int]]:
    if not spans:
        return [(texto, mapa[0] if mapa else 0)]
    segmentos: list[tuple[str, int]] = []
    pos = 0
    for inicio, fin, reemplazo, _ in sorted(spans, key=lambda s: s[0]):
        if pos < inicio:
            fmt = mapa[pos] if pos < len(mapa) else (mapa[-1] if mapa else 0)
            segmentos.append((texto[pos:inicio], fmt))
        fmt_rep = mapa[inicio] if inicio < len(mapa) else (mapa[-1] if mapa else 0)
        segmentos.append((reemplazo, fmt_rep))
        pos = fin
    if pos < len(texto):
        fmt = mapa[pos] if pos < len(mapa) else (mapa[-1] if mapa else 0)
        segmentos.append((texto[pos:], fmt))
    return segmentos


def _fusionar_segmentos(segmentos: list[tuple[str, int]]) -> list[tuple[str, int]]:
    if not segmentos:
        return []
    fusionados = [segmentos[0]]
    for texto, fmt in segmentos[1:]:
        if not texto:
            continue
        if fmt == fusionados[-1][1]:
            fusionados[-1] = (fusionados[-1][0] + texto, fmt)
        else:
            fusionados.append((texto, fmt))
    return fusionados


def _eliminar_runs(p) -> None:
    for run in list(p.runs):
        run._element.getparent().remove(run._element)


def _aplicar_a_parrafo(p, motor: MotorAnonimizacion, hallazgos: list[Hallazgo],
                       registro: dict):
    """Aplica reemplazos conservando el formato (subrayado, negrita, etc.) de cada run."""
    texto_original = p.text
    if not texto_original:
        return

    spans = motor._recolectar_spans(texto_original, hallazgos)
    if not spans:
        return

    mapa = _mapear_caracteres_runs(p)
    rpr_copias = [_copiar_rpr(run) for run in p.runs]
    segmentos = _fusionar_segmentos(_segmentos_con_formato(texto_original, mapa, spans))

    _eliminar_runs(p)
    for texto, fmt_idx in segmentos:
        if not texto:
            continue
        nuevo_run = p.add_run(texto)
        if rpr_copias and fmt_idx < len(rpr_copias):
            _aplicar_rpr(nuevo_run, rpr_copias[fmt_idx])

    for _, _, reemplazo, original in spans:
        registro[original] = reemplazo


def _procesar_contenedor(contenedor, motor: MotorAnonimizacion,
                         hallazgos: list[Hallazgo], registro: dict):
    for p in contenedor.paragraphs:
        _aplicar_a_parrafo(p, motor, hallazgos, registro)
    for tabla in contenedor.tables:
        for fila in tabla.rows:
            for celda in fila.cells:
                _procesar_contenedor(celda, motor, hallazgos, registro)


def _anonimizar_xml_wt(xml_texto: str, motor: MotorAnonimizacion,
                       hallazgos: list[Hallazgo], registro: dict) -> str:
    def _sub(m):
        contenido = unescape(m.group(2))
        nuevo = motor.aplicar(contenido, hallazgos, registro)
        return m.group(1) + escape(nuevo) + m.group(3)

    return re.sub(r"(<w:t[^>]*>)(.*?)(</w:t>)", _sub, xml_texto, flags=re.DOTALL)


def _postprocesar_zip(ruta_docx: Path, motor: MotorAnonimizacion,
                      hallazgos: list[Hallazgo], registro: dict):
    objetivos = {"word/footnotes.xml", "word/endnotes.xml", "word/comments.xml"}
    fd, tmp_nombre = tempfile.mkstemp(suffix=".docx")
    tmp = Path(tmp_nombre)
    os.close(fd)
    with zipfile.ZipFile(ruta_docx, "r") as zin, \
         zipfile.ZipFile(tmp, "w", zipfile.ZIP_DEFLATED) as zout:
        for item in zin.infolist():
            datos = zin.read(item.filename)
            if item.filename in objetivos:
                xml = datos.decode("utf-8")
                datos = _anonimizar_xml_wt(xml, motor, hallazgos, registro).encode("utf-8")
            zout.writestr(item, datos)
    shutil.move(str(tmp), str(ruta_docx))


def _carpeta_destino(ruta: Path, carpeta_salida: Path | None) -> Path:
    destino = carpeta_salida if carpeta_salida else ruta.parent
    destino.mkdir(parents=True, exist_ok=True)
    return destino


def procesar_docx(ruta: Path, motor: MotorAnonimizacion,
                  hallazgos: list[Hallazgo], registro: dict,
                  carpeta_salida: Path | None = None) -> Path:
    from docx import Document
    doc = Document(ruta)
    _procesar_contenedor(doc, motor, hallazgos, registro)
    for seccion in doc.sections:
        for parte in (seccion.header, seccion.footer,
                      seccion.first_page_header, seccion.first_page_footer,
                      seccion.even_page_header, seccion.even_page_footer):
            if parte is not None:
                _procesar_contenedor(parte, motor, hallazgos, registro)
    cp = doc.core_properties
    cp.author = cp.last_modified_by = cp.title = cp.subject = ""
    cp.comments = cp.keywords = cp.category = ""
    salida = _carpeta_destino(ruta, carpeta_salida) / f"{ruta.stem}_ANONIMIZADO.docx"
    doc.save(salida)
    _postprocesar_zip(salida, motor, hallazgos, registro)
    return salida


def procesar_txt(ruta: Path, motor: MotorAnonimizacion,
                 hallazgos: list[Hallazgo], registro: dict,
                 carpeta_salida: Path | None = None) -> Path:
    texto = ruta.read_text(encoding="utf-8", errors="ignore")
    salida = _carpeta_destino(ruta, carpeta_salida) / f"{ruta.stem}_ANONIMIZADO{ruta.suffix}"
    salida.write_text(motor.aplicar(texto, hallazgos, registro), encoding="utf-8")
    return salida


def guardar_equivalencias(ruta: Path, registro: dict,
                          carpeta_salida: Path | None = None) -> Path:
    tabla = _carpeta_destino(ruta, carpeta_salida) / f"{ruta.stem}_EQUIVALENCIAS.csv"
    with open(tabla, "w", newline="", encoding="utf-8-sig") as f:
        w = csv.writer(f)
        w.writerow(["Dato real", "Reemplazo", "Tipo"])
        for real, rol in sorted(registro.items(), key=lambda x: x[0].lower()):
            w.writerow([real, rol, ""])
    return tabla


def procesar(ruta: Path, hallazgos: list[Hallazgo],
             carpeta_salida: Path | None = None):
    motor = MotorAnonimizacion()
    registro: dict[str, str] = {}
    suf = ruta.suffix.lower()
    if suf == ".docx":
        salida = procesar_docx(ruta, motor, hallazgos, registro, carpeta_salida)
    elif suf in (".txt", ".md"):
        salida = procesar_txt(ruta, motor, hallazgos, registro, carpeta_salida)
    elif suf == ".pdf":
        salida = procesar_pdf(ruta, motor, hallazgos, registro, carpeta_salida)
    else:
        raise ValueError("Formato no soportado. Usa .docx, .txt, .md o .pdf")
    tabla = guardar_equivalencias(ruta, registro, carpeta_salida)
    parrafos_salida = extraer_parrafos(salida)
    fugas = motor.verificar_resultado(parrafos_salida, hallazgos)
    return salida, tabla, len(registro), fugas


def _parsear_manual(texto_caja: str) -> dict:
    nuevos = {}
    for linea in texto_caja.splitlines():
        linea = linea.strip()
        if not linea or linea.startswith("#"):
            continue
        if "=>" in linea:
            real, rol = linea.split("=>", 1)
            if real.strip() and rol.strip():
                nuevos[real.strip()] = rol.strip()
    return nuevos


def _crear_tabla_revision(padre, items: list[Hallazgo], titulo_cols: tuple,
                          valores_defecto: bool):
    import tkinter as tk
    from tkinter import ttk

    marco = tk.Frame(padre, bg="#111111")
    marco.pack(fill="both", expand=True)

    canvas = tk.Canvas(marco, bg="#1d1d1d", highlightthickness=0)
    scroll = ttk.Scrollbar(marco, orient="vertical", command=canvas.yview)
    interior = tk.Frame(canvas, bg="#1d1d1d")
    interior.bind("<Configure>",
                  lambda e: canvas.configure(scrollregion=canvas.bbox("all")))
    canvas.create_window((0, 0), window=interior, anchor="nw")
    canvas.configure(yscrollcommand=scroll.set)
    canvas.pack(side="left", fill="both", expand=True)
    scroll.pack(side="right", fill="y")

    vars_check: list[tk.BooleanVar] = []
    colores_tipo = {
        "PERSONA": "#7ec8ff", "EMPRESA": "#ffb347", "LUGAR": "#b19cd9",
        "CÉDULA": "#ff6b6b", "NIT": "#ff6b6b", "CORREO": "#ff6b6b",
        "TELÉFONO": "#ff6b6b", "DIRECCIÓN": "#ff6b6b", "RADICADO": "#aaaaaa",
        "MANUAL": "#9be29b", "REVISAR": "#ffcc66",
    }

    encabezado = tk.Frame(interior, bg="#2a2a2a")
    encabezado.pack(fill="x", pady=(0, 4))
    for i, (txt, w) in enumerate(titulo_cols):
        tk.Label(encabezado, text=txt, bg="#2a2a2a", fg="#cccccc",
                 font=("Segoe UI", 9, "bold"), width=w, anchor="w"
                 ).grid(row=0, column=i, padx=4, pady=4, sticky="w")

    for h in items:
        fila = tk.Frame(interior, bg="#1d1d1d")
        fila.pack(fill="x", pady=1)
        var = tk.BooleanVar(value=valores_defecto if h.tipo == "REVISAR" else h.activo)
        vars_check.append(var)
        color = colores_tipo.get(h.tipo, "#f5f5f5")
        tk.Checkbutton(fila, variable=var, bg="#1d1d1d", activebackground="#1d1d1d",
                       selectcolor="#333333").grid(row=0, column=0, padx=4)
        tk.Label(fila, text=h.texto[:45] + ("…" if len(h.texto) > 45 else ""),
                 bg="#1d1d1d", fg="#f5f5f5", font=("Consolas", 9),
                 width=30, anchor="w").grid(row=0, column=1, sticky="w")
        tk.Label(fila, text=h.reemplazo, bg="#1d1d1d", fg="#9be29b",
                 font=("Consolas", 9), width=18, anchor="w").grid(row=0, column=2, sticky="w")
        tk.Label(fila, text=h.tipo, bg="#1d1d1d", fg=color,
                 font=("Segoe UI", 8), width=10, anchor="w").grid(row=0, column=3, sticky="w")
        tk.Label(fila, text=h.fuente, bg="#1d1d1d", fg="#888888",
                 font=("Segoe UI", 8), width=10, anchor="w").grid(row=0, column=4, sticky="w")
        conf_color = {"alta": "#9be29b", "media": "#ffcc66", "baja": "#ff6b6b"}.get(h.confianza, "#888")
        tk.Label(fila, text=h.confianza, bg="#1d1d1d", fg=conf_color,
                 font=("Segoe UI", 8), width=8, anchor="w").grid(row=0, column=5, sticky="w")

    return vars_check


def mostrar_revision(raiz, hallazgos: list[Hallazgo], texto: str,
                     motor: MotorAnonimizacion,
                     parrafos: list[str] | None = None) -> list[Hallazgo] | None:
    import tkinter as tk
    from tkinter import ttk

    resultado: list[Hallazgo] | None = None
    bloques = parrafos if parrafos else texto.split("\n")
    pendientes = motor.detectar_pendientes(texto, hallazgos, parrafos=bloques)

    ventana = tk.Toplevel(raiz)
    ventana.title("Revisión — confirma qué anonimizar")
    ventana.geometry("780x560")
    ventana.configure(bg="#111111")
    ventana.transient(raiz)
    ventana.grab_set()

    tk.Label(
        ventana,
        text=(f"Detectados: {len(hallazgos)} · Pendientes por revisar: {len(pendientes)}. "
              "Revisa la pestaña «Pendientes» antes de confirmar."),
        bg="#111111", fg="#f5f5f5", font=("Segoe UI", 10), wraplength=720,
    ).pack(pady=(12, 8), padx=16)

    notebook = ttk.Notebook(ventana)
    notebook.pack(fill="both", expand=True, padx=16, pady=4)

    tab_detectados = tk.Frame(notebook, bg="#111111")
    tab_pendientes = tk.Frame(notebook, bg="#111111")
    notebook.add(tab_detectados, text=f"Detectados ({len(hallazgos)})")
    notebook.add(tab_pendientes, text=f"Pendientes ({len(pendientes)})")

    cols = [("✓", 3), ("Dato", 30), ("Reemplazo", 18), ("Tipo", 10), ("Fuente", 10), ("Conf.", 8)]
    vars_hallazgos = _crear_tabla_revision(tab_detectados, hallazgos, cols, True)
    vars_pendientes = _crear_tabla_revision(tab_pendientes, pendientes, cols, False) if pendientes else []

    if not pendientes:
        tk.Label(tab_pendientes, text="No quedaron números ni nombres sospechosos sin anonimizar.",
                 bg="#111111", fg="#9be29b", font=("Segoe UI", 10)).pack(pady=40)

    def confirmar():
        nonlocal resultado
        for h, var in zip(hallazgos, vars_hallazgos):
            h.activo = var.get()
        extras: list[Hallazgo] = []
        for h, var in zip(pendientes, vars_pendientes):
            if var.get():
                h.activo = True
                h.tipo = "MANUAL"
                h.fuente = "revisión"
                h.patron = re.compile(r"(?<!\w)" + re.escape(h.texto) + r"(?!\w)", re.IGNORECASE)
                extras.append(h)
        resultado = hallazgos + extras
        ventana.destroy()

    def cancelar():
        ventana.destroy()

    marco_btn = tk.Frame(ventana, bg="#111111")
    marco_btn.pack(pady=12)
    tk.Button(marco_btn, text="Confirmar y anonimizar", command=confirmar,
              bg="#ffffff", fg="#111111", font=("Segoe UI", 10, "bold"),
              relief="flat", padx=16, pady=6).pack(side="left", padx=8)
    tk.Button(marco_btn, text="Cancelar", command=cancelar,
              bg="#333333", fg="#f5f5f5", font=("Segoe UI", 10),
              relief="flat", padx=16, pady=6).pack(side="left", padx=8)

    ventana.wait_window()
    return resultado


def _pestaña_revision_doc(padre, hallazgos: list[Hallazgo], pendientes: list[Hallazgo]):
    import tkinter as tk
    from tkinter import ttk

    notebook = ttk.Notebook(padre)
    notebook.pack(fill="both", expand=True)

    tab_detectados = tk.Frame(notebook, bg="#111111")
    tab_pendientes = tk.Frame(notebook, bg="#111111")
    notebook.add(tab_detectados, text=f"Detectados ({len(hallazgos)})")
    notebook.add(tab_pendientes, text=f"Pendientes ({len(pendientes)})")

    cols = [("✓", 3), ("Dato", 30), ("Reemplazo", 18), ("Tipo", 10), ("Fuente", 10), ("Conf.", 8)]
    vars_hallazgos = _crear_tabla_revision(tab_detectados, hallazgos, cols, True)
    vars_pendientes = (
        _crear_tabla_revision(tab_pendientes, pendientes, cols, False) if pendientes else []
    )

    if not pendientes:
        tk.Label(
            tab_pendientes,
            text="No quedaron números ni nombres sospechosos sin anonimizar.",
            bg="#111111", fg="#9be29b", font=("Segoe UI", 10),
        ).pack(pady=40)

    return vars_hallazgos, vars_pendientes, pendientes


def mostrar_revision_multiples(raiz, documentos: list[dict]) -> list[dict] | None:
    import tkinter as tk
    from tkinter import ttk

    resultado: list[dict] | None = None
    total_h = sum(len(d["hallazgos"]) for d in documentos)
    total_p = sum(
        len(d["motor"].detectar_pendientes(d["texto"], d["hallazgos"], parrafos=d["parrafos"]))
        for d in documentos
    )

    ventana = tk.Toplevel(raiz)
    ventana.title("Revisión por documento")
    ventana.geometry("820x600")
    ventana.configure(bg="#111111")
    ventana.transient(raiz)
    ventana.grab_set()

    tk.Label(
        ventana,
        text=(f"{len(documentos)} documento(s) · {total_h} detectados · "
              f"{total_p} pendientes. Revisa cada pestaña antes de confirmar."),
        bg="#111111", fg="#f5f5f5", font=("Segoe UI", 10), wraplength=760,
    ).pack(pady=(12, 8), padx=16)

    notebook = ttk.Notebook(ventana)
    notebook.pack(fill="both", expand=True, padx=16, pady=4)

    estados: list[tuple] = []
    for doc in documentos:
        pendientes = doc["motor"].detectar_pendientes(
            doc["texto"], doc["hallazgos"], parrafos=doc["parrafos"])
        tab = tk.Frame(notebook, bg="#111111")
        nombre = doc["ruta"].name
        if len(nombre) > 28:
            nombre = nombre[:25] + "…"
        notebook.add(tab, text=f"{nombre} ({len(doc['hallazgos'])})")
        vars_h, vars_p, pend = _pestaña_revision_doc(tab, doc["hallazgos"], pendientes)
        estados.append((doc, vars_h, vars_p, pend))

    def confirmar():
        nonlocal resultado
        salida: list[dict] = []
        for doc, vars_h, vars_p, pendientes in estados:
            hallazgos = doc["hallazgos"]
            for h, var in zip(hallazgos, vars_h):
                h.activo = var.get()
            extras: list[Hallazgo] = []
            for h, var in zip(pendientes, vars_p):
                if var.get():
                    h.activo = True
                    h.tipo = "MANUAL"
                    h.fuente = "revisión"
                    h.patron = re.compile(
                        r"(?<!\w)" + re.escape(h.texto) + r"(?!\w)", re.IGNORECASE)
                    extras.append(h)
            salida.append({**doc, "hallazgos": hallazgos + extras})
        resultado = salida
        ventana.destroy()

    def cancelar():
        ventana.destroy()

    marco_btn = tk.Frame(ventana, bg="#111111")
    marco_btn.pack(pady=12)
    tk.Button(marco_btn, text="Confirmar y anonimizar todo", command=confirmar,
              bg="#ffffff", fg="#111111", font=("Segoe UI", 10, "bold"),
              relief="flat", padx=16, pady=6).pack(side="left", padx=8)
    tk.Button(marco_btn, text="Cancelar", command=cancelar,
              bg="#333333", fg="#f5f5f5", font=("Segoe UI", 10),
              relief="flat", padx=16, pady=6).pack(side="left", padx=8)

    ventana.wait_window()
    return resultado


def _carpeta_sugerida(archivos: list[Path]) -> Path:
    prefs = cargar_preferencias()
    if prefs.get("modo_salida") == "carpeta":
        guardada = prefs.get("carpeta_salida", "").strip()
        if guardada:
            p = Path(guardada)
            if p.is_dir():
                return p.resolve()
    if archivos:
        padres = {p.parent.resolve() for p in archivos}
        if len(padres) == 1:
            return next(iter(padres))
        return archivos[0].parent.resolve()
    return Path.home() / "Documents"


def mostrar_dialogo_carpeta_salida(raiz, archivos: list[Path]) -> tuple[bool, Path | None]:
    """(True, None) junto al original · (True, Path) carpeta única · (False, None) cancelado."""
    import tkinter as tk
    from tkinter import filedialog, messagebox

    prefs = cargar_preferencias()
    modo_inicial = prefs.get("modo_salida", "junto")
    carpeta_inicial = _carpeta_sugerida(archivos)

    ventana = tk.Toplevel(raiz)
    ventana.title("¿Dónde guardar los archivos?")
    ventana.geometry("560x280")
    ventana.configure(bg="#111111")
    ventana.transient(raiz)
    ventana.grab_set()

    resultado: dict = {"ok": False, "carpeta": None}

    tk.Label(
        ventana,
        text="Elige la carpeta de destino antes de generar los documentos anonimizados.",
        bg="#111111", fg="#f5f5f5", font=("Segoe UI", 10), wraplength=500,
    ).pack(pady=(16, 12), padx=20)

    marco = tk.Frame(ventana, bg="#111111")
    marco.pack(fill="x", padx=20)

    var_modo = tk.StringVar(value=modo_inicial)
    tk.Radiobutton(
        marco, text="Junto a cada archivo original", variable=var_modo, value="junto",
        bg="#111111", fg="#f5f5f5", selectcolor="#333333", activebackground="#111111",
        font=("Segoe UI", 10),
    ).pack(anchor="w", pady=2)
    tk.Radiobutton(
        marco, text="Carpeta única para todos:", variable=var_modo, value="carpeta",
        bg="#111111", fg="#f5f5f5", selectcolor="#333333", activebackground="#111111",
        font=("Segoe UI", 10),
    ).pack(anchor="w", pady=(8, 4))

    fila_ruta = tk.Frame(marco, bg="#111111")
    fila_ruta.pack(fill="x", pady=4)
    entrada = tk.Entry(fila_ruta, font=("Consolas", 9), bg="#1d1d1d", fg="#f5f5f5",
                     insertbackground="#ffffff", relief="flat")
    entrada.insert(0, str(carpeta_inicial))
    entrada.pack(side="left", fill="x", expand=True, ipady=4)

    def examinar():
        elegida = filedialog.askdirectory(
            title="Carpeta de salida",
            initialdir=entrada.get().strip() or str(Path.home()),
        )
        if elegida:
            entrada.delete(0, "end")
            entrada.insert(0, elegida)

    tk.Button(fila_ruta, text="Examinar…", command=examinar,
              bg="#333333", fg="#f5f5f5", font=("Segoe UI", 9),
              relief="flat", padx=10, pady=4).pack(side="left", padx=(8, 0))

    var_recordar = tk.BooleanVar(value=bool(prefs.get("modo_salida")))
    tk.Checkbutton(
        ventana, text="Usar esta opción como predeterminada",
        variable=var_recordar, bg="#111111", fg="#aaaaaa", selectcolor="#333333",
        activebackground="#111111", font=("Segoe UI", 9),
    ).pack(anchor="w", padx=20, pady=(12, 0))

    def confirmar():
        if var_modo.get() == "carpeta":
            texto = entrada.get().strip()
            if not texto:
                messagebox.showwarning("Falta carpeta", "Indica o selecciona una carpeta.")
                return
            destino = Path(texto)
            try:
                destino.mkdir(parents=True, exist_ok=True)
            except OSError as e:
                messagebox.showerror("Carpeta no válida", str(e))
                return
            if not destino.is_dir():
                messagebox.showwarning("Carpeta no válida", "La ruta no es una carpeta.")
                return
            resultado["carpeta"] = destino.resolve()
        else:
            resultado["carpeta"] = None
        if var_recordar.get():
            if var_modo.get() == "carpeta":
                guardar_preferencias({
                    "modo_salida": "carpeta",
                    "carpeta_salida": str(resultado["carpeta"]),
                })
            else:
                guardar_preferencias({"modo_salida": "junto", "carpeta_salida": ""})
        resultado["ok"] = True
        ventana.destroy()

    def cancelar():
        ventana.destroy()

    marco_btn = tk.Frame(ventana, bg="#111111")
    marco_btn.pack(pady=20)
    tk.Button(marco_btn, text="Continuar y anonimizar", command=confirmar,
              bg="#ffffff", fg="#111111", font=("Segoe UI", 10, "bold"),
              relief="flat", padx=16, pady=6).pack(side="left", padx=8)
    tk.Button(marco_btn, text="Cancelar", command=cancelar,
              bg="#333333", fg="#f5f5f5", font=("Segoe UI", 10),
              relief="flat", padx=16, pady=6).pack(side="left", padx=8)

    ventana.wait_window()
    return resultado["ok"], resultado["carpeta"]


def _texto_salida_predeterminada() -> str:
    prefs = cargar_preferencias()
    if prefs.get("modo_salida") == "carpeta":
        carpeta = prefs.get("carpeta_salida", "").strip()
        if carpeta:
            return f"Salida predeterminada: {carpeta}"
    return "Salida predeterminada: junto a cada archivo original"


def lanzar_gui():
    import tkinter as tk
    from tkinter import filedialog, messagebox, scrolledtext

    reemplazos = cargar_reemplazos()
    archivos_sel: list[Path] = []
    analisis_docs: list[dict] = []

    raiz = tk.Tk()
    raiz.title("Anonimizador de Expedientes — Trifuerza v7")
    raiz.geometry("720x660")
    raiz.configure(bg="#111111")
    estilo = {"bg": "#111111", "fg": "#f5f5f5", "font": ("Segoe UI", 10)}

    tk.Label(raiz, text="ANONIMIZADOR DE EXPEDIENTES",
             font=("Segoe UI", 14, "bold"), bg="#111111", fg="#ffffff"
             ).pack(pady=(16, 2))
    tk.Label(raiz, text="Varios archivos · .docx .txt .md .pdf · Revisión local",
             **estilo).pack(pady=(0, 12))

    marco_archivo = tk.Frame(raiz, bg="#111111")
    marco_archivo.pack(fill="x", padx=20)
    etiqueta_archivo = tk.Label(marco_archivo, text="Ningún archivo seleccionado",
                                anchor="w", **estilo)

    def _actualizar_etiqueta():
        if not archivos_sel:
            etiqueta_archivo.config(text="Ningún archivo seleccionado")
        elif len(archivos_sel) == 1:
            etiqueta_archivo.config(text=archivos_sel[0].name)
        else:
            nombres = ", ".join(p.name for p in archivos_sel[:3])
            if len(archivos_sel) > 3:
                nombres += f" … (+{len(archivos_sel) - 3})"
            etiqueta_archivo.config(text=f"{len(archivos_sel)} archivos: {nombres}")

    def seleccionar():
        rutas = filedialog.askopenfilenames(
            title="Selecciona uno o varios documentos",
            filetypes=FILTRO_ARCHIVOS)
        if rutas:
            archivos_sel.clear()
            analisis_docs.clear()
            for r in rutas:
                p = Path(r)
                if es_formato_soportado(p):
                    archivos_sel.append(p)
            _actualizar_etiqueta()
            log(f"{len(archivos_sel)} archivo(s) seleccionado(s). Pulsa ANALIZAR.")

    def seleccionar_carpeta():
        carpeta = filedialog.askdirectory(title="Selecciona una carpeta con documentos")
        if carpeta:
            try:
                encontrados = recolectar_rutas([carpeta])
            except ValueError as e:
                messagebox.showwarning("Sin archivos válidos", str(e))
                return
            if not encontrados:
                messagebox.showwarning("Carpeta vacía",
                                       "No hay .docx, .txt, .md ni .pdf en esa carpeta.")
                return
            archivos_sel.clear()
            analisis_docs.clear()
            archivos_sel.extend(encontrados)
            _actualizar_etiqueta()
            log(f"{len(archivos_sel)} archivo(s) en carpeta. Pulsa ANALIZAR.")

    tk.Button(marco_archivo, text="Seleccionar archivos…", command=seleccionar,
              bg="#ffffff", fg="#111111", font=("Segoe UI", 10, "bold"),
              relief="flat", padx=12, pady=4).pack(side="left")
    tk.Button(marco_archivo, text="Carpeta…", command=seleccionar_carpeta,
              bg="#333333", fg="#f5f5f5", font=("Segoe UI", 10),
              relief="flat", padx=10, pady=4).pack(side="left", padx=(6, 0))
    etiqueta_archivo.pack(side="left", padx=12)

    etiqueta_salida = tk.Label(raiz, text=_texto_salida_predeterminada(),
                               anchor="w", fg="#888888", bg="#111111",
                               font=("Segoe UI", 9))
    etiqueta_salida.pack(fill="x", padx=20, pady=(4, 0))

    tk.Label(raiz, text="Reemplazos adicionales (opcional, formato:  Nombre => rol):",
             anchor="w", **estilo).pack(fill="x", padx=20, pady=(12, 2))
    tk.Label(raiz,
             text="Cédulas, NIT, nombres y empresas se detectan solos. "
                  "Cortes, leyes y juzgados NO se anonimizan.",
             anchor="w", fg="#aaaaaa", bg="#111111", font=("Segoe UI", 9)
             ).pack(fill="x", padx=20, pady=(0, 4))

    caja = scrolledtext.ScrolledText(raiz, height=5, font=("Consolas", 10),
                                     bg="#1d1d1d", fg="#f5f5f5",
                                     insertbackground="#ffffff", relief="flat")
    caja.pack(fill="both", expand=False, padx=20)
    if reemplazos:
        caja.insert("1.0", "\n".join(f"{k} => {v}" for k, v in reemplazos.items()))

    salida_log = scrolledtext.ScrolledText(raiz, height=10, font=("Consolas", 9),
                                           bg="#1d1d1d", fg="#9be29b",
                                           relief="flat", state="disabled")

    def log(msg):
        salida_log.config(state="normal")
        salida_log.insert("end", msg + "\n")
        salida_log.config(state="disabled")
        salida_log.see("end")

    def analizar():
        if not archivos_sel:
            messagebox.showwarning("Faltan archivos", "Primero selecciona documento(s) o carpeta.")
            return
        manual = _parsear_manual(caja.get("1.0", "end"))
        guardar_reemplazos(manual)
        analisis_docs.clear()
        errores: list[str] = []
        try:
            for ruta in archivos_sel:
                try:
                    texto = extraer_texto(ruta)
                    parrafos = extraer_parrafos(ruta)
                    motor = MotorAnonimizacion(reemplazos_manual=manual)
                    hallazgos = motor.analizar(texto)
                    analisis_docs.append({
                        "ruta": ruta,
                        "texto": texto,
                        "parrafos": parrafos,
                        "motor": motor,
                        "hallazgos": hallazgos,
                    })
                    log(f"✔ {ruta.name}: {len(hallazgos)} hallazgo(s)")
                except Exception as e:
                    errores.append(f"{ruta.name}: {e}")
                    log(f"✗ {ruta.name}: {e}")
            if not analisis_docs:
                messagebox.showerror("Error", "\n".join(errores) or "No se pudo analizar.")
                return
            total = sum(len(d["hallazgos"]) for d in analisis_docs)
            log(f"Análisis listo: {len(analisis_docs)} doc(s), {total} hallazgo(s) en total.")
            if errores:
                log(f"⚠ {len(errores)} archivo(s) con error (omitidos).")
            log("Pulsa ANONIMIZAR para revisar y confirmar.")
        except ImportError as e:
            messagebox.showerror("Falta una librería", str(e))

    def anonimizar():
        if not archivos_sel:
            messagebox.showwarning("Faltan archivos", "Primero selecciona documento(s) o carpeta.")
            return
        if not analisis_docs:
            analizar()
            if not analisis_docs:
                return
        if len(analisis_docs) == 1:
            doc = analisis_docs[0]
            confirmados = mostrar_revision(
                raiz, list(doc["hallazgos"]), doc["texto"], doc["motor"],
                parrafos=doc["parrafos"])
            docs_ok = [{**doc, "hallazgos": confirmados}] if confirmados else None
        else:
            docs_ok = mostrar_revision_multiples(raiz, analisis_docs)
        if not docs_ok:
            log("Anonimización cancelada.")
            return
        ok_salida, carpeta_dest = mostrar_dialogo_carpeta_salida(
            raiz, [doc["ruta"] for doc in docs_ok])
        if not ok_salida:
            log("Anonimización cancelada (sin carpeta de salida).")
            return
        etiqueta_salida.config(text=_texto_salida_predeterminada())
        if carpeta_dest:
            log(f"Guardando en: {carpeta_dest}")
        else:
            log("Guardando junto a cada archivo original.")
        procesados = 0
        total_fugas = 0
        carpeta_mostrar = carpeta_dest
        try:
            for doc in docs_ok:
                activos = [h for h in doc["hallazgos"] if h.activo]
                if not activos:
                    log(f"⊘ {doc['ruta'].name}: sin datos marcados, omitido.")
                    continue
                salida, tabla, n, fugas = procesar(
                    doc["ruta"], doc["hallazgos"], carpeta_salida=carpeta_dest)
                procesados += 1
                total_fugas += len(fugas)
                if carpeta_mostrar is None:
                    carpeta_mostrar = salida.parent
                log(f"✔ {salida} ({n} reemplazos)")
                if fugas:
                    for f in fugas[:3]:
                        log(f"  ⚠ fuga: {f.texto}")
            if procesados == 0:
                messagebox.showwarning("Sin cambios", "Ningún documento tenía datos marcados.")
                return
            log(f"⚠ {ADVERTENCIA}")
            msg = f"{procesados} documento(s) anonimizado(s)."
            if carpeta_dest:
                msg += f"\n\nCarpeta:\n{carpeta_dest}"
            if total_fugas:
                msg += f"\n\n⚠ Quedaron {total_fugas} dato(s) visibles en total."
            messagebox.showinfo("Listo", msg)
        except Exception as e:
            messagebox.showerror("Error", str(e))

    marco_btn = tk.Frame(raiz, bg="#111111")
    marco_btn.pack(pady=10)
    tk.Button(marco_btn, text="ANALIZAR", command=analizar,
              bg="#333333", fg="#f5f5f5", font=("Segoe UI", 10, "bold"),
              relief="flat", padx=20, pady=6).pack(side="left", padx=6)
    tk.Button(marco_btn, text="ANONIMIZAR", command=anonimizar,
              bg="#ffffff", fg="#111111", font=("Segoe UI", 11, "bold"),
              relief="flat", padx=24, pady=8).pack(side="left", padx=6)

    salida_log.pack(fill="both", expand=True, padx=20, pady=(0, 16))
    log("Selecciona archivos o una carpeta y pulsa ANALIZAR.")
    raiz.mainloop()


def _parsear_args_cli(argv: list[str]) -> tuple[list[str], Path | None]:
    args = list(argv)
    carpeta: Path | None = None
    i = 0
    while i < len(args):
        if args[i] == "--salida":
            if i + 1 >= len(args):
                print("Error: --salida requiere una ruta de carpeta.")
                sys.exit(1)
            carpeta = Path(args[i + 1])
            del args[i:i + 2]
            continue
        i += 1
    return args, carpeta


def _resolver_carpeta_salida_cli(carpeta_arg: Path | None) -> Path | None:
    if carpeta_arg is not None:
        carpeta_arg.mkdir(parents=True, exist_ok=True)
        if not carpeta_arg.is_dir():
            print(f"Error: --salida no es una carpeta válida: {carpeta_arg}")
            sys.exit(1)
        return carpeta_arg.resolve()
    prefs = cargar_preferencias()
    if prefs.get("modo_salida") == "carpeta":
        guardada = prefs.get("carpeta_salida", "").strip()
        if guardada:
            p = Path(guardada)
            if p.is_dir():
                return p.resolve()
    return None


def main():
    if len(sys.argv) >= 2:
        entradas, carpeta_arg = _parsear_args_cli(sys.argv[1:])
        if not entradas:
            print("Uso: python anonimizador.py archivo.docx [más archivos o carpeta/] [--salida carpeta/]")
            sys.exit(1)
        try:
            rutas = recolectar_rutas(entradas)
        except (FileNotFoundError, ValueError) as e:
            print(f"Error: {e}")
            sys.exit(1)
        if not rutas:
            print("No hay archivos .docx, .txt, .md ni .pdf para procesar.")
            sys.exit(1)
        carpeta_salida = _resolver_carpeta_salida_cli(carpeta_arg)
        manual = cargar_reemplazos()
        errores = 0
        for ruta in rutas:
            try:
                texto = extraer_texto(ruta)
                motor = MotorAnonimizacion(reemplazos_manual=manual)
                hallazgos = motor.analizar(texto)
                salida, tabla, n, fugas = procesar(ruta, hallazgos, carpeta_salida=carpeta_salida)
                print(f"✔ {ruta.name} → {salida.name}")
                print(f"  Equivalencias: {tabla.name}  (NUNCA subas este archivo)")
                print(f"  Reemplazos: {n}")
                if fugas:
                    print(f"  ⚠ Quedaron {len(fugas)} dato(s) sin anonimizar")
            except Exception as e:
                errores += 1
                print(f"✗ {ruta.name}: {e}")
        if errores:
            sys.exit(1)
        print(f"⚠ {ADVERTENCIA}")
    else:
        lanzar_gui()


if __name__ == "__main__":
    main()
